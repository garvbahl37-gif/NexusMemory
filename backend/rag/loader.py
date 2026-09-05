"""Document loading.

PDFs are read with pypdf directly rather than through a loader wrapper, which
keeps page numbers (used for citations) explicit and drops a large dependency.
"""
import csv
import logging
from pathlib import Path

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def _load_pdf(file_path: str, name: str) -> list[Document]:
    """One Document per page, so citations can name the page they came from."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for number, page in enumerate(reader.pages):
        content = (page.extract_text() or "").strip()
        if content:
            pages.append(
                Document(
                    page_content=content,
                    metadata={"source": name, "page": number},
                )
            )

    logger.info(f"Loaded PDF: {name} — {len(pages)} pages with text")
    return pages


def _load_text(file_path: str, name: str) -> list[Document]:
    content = Path(file_path).read_text(encoding="utf-8", errors="ignore")
    logger.info(f"Loaded text file: {name}")
    return [Document(page_content=content, metadata={"source": name, "page": 0})]


def _load_csv(file_path: str, name: str) -> list[Document]:
    """Read a CSV into a single text Document (header-aware)."""
    rows = []
    with open(file_path, newline="", encoding="utf-8", errors="ignore") as f:
        for row in csv.reader(f):
            rows.append(", ".join(cell.strip() for cell in row))
    logger.info(f"Loaded CSV: {name} — {len(rows)} rows")
    return [
        Document(
            page_content="\n".join(rows),
            metadata={"source": name, "page": 0},
        )
    ]


def _load_docx(file_path: str, name: str) -> list[Document]:
    """Extract paragraphs and table cells from a .docx."""
    from docx import Document as DocxDocument  # python-docx

    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    logger.info(f"Loaded DOCX: {name} — {len(parts)} blocks")
    return [
        Document(
            page_content="\n".join(parts),
            metadata={"source": name, "page": 0},
        )
    ]


LOADERS = {
    ".pdf": _load_pdf,
    ".txt": _load_text,
    ".md": _load_text,
    ".csv": _load_csv,
    ".docx": _load_docx,
}


def load_document(file_path: str) -> list[Document]:
    """Load a document by extension. Supports PDF, TXT, MD, DOCX and CSV."""
    path = Path(file_path)
    loader = LOADERS.get(path.suffix.lower())

    if loader is None:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    try:
        return loader(file_path, path.name)
    except Exception as e:
        logger.error(f"Failed to load {path.name}: {e}")
        raise
